import json
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches
from plotly.utils import PlotlyJSONEncoder
from typing import Optional
from pydantic import BaseModel
from fastapi import FastAPI, UploadFile, File, Form
from fastapi.responses import FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi import HTTPException
from utils import filter_data, generate_chart, generate_insights, generate_correlation_matrix

app = FastAPI()

# Set CORS policies
origins = [
    "http://localhost",
    "http://localhost:3000",
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.post("/generate_dashboard")
async def generate_dashboard(
    x_column: str = Form(...),
    y_column: str = Form(...),
    chart_type: str = Form(...),
    filter_column: Optional[str] = Form(''),
    time_range_start: Optional[str] = Form(''),
    time_range_end: Optional[str] = Form(''),
    file: UploadFile = File(...)
):
    # Read and preprocess the CSV data
    df = pd.read_csv(file.file)

    # Filter the data based on the time range, if filter_column is provided
    if filter_column and time_range_start and time_range_end:
        df = filter_data(df, filter_column, time_range_start, time_range_end)

    # Generate the chart
    fig = generate_chart(df, x_column, y_column, chart_type)

    # Convert the chart to JSON
    chart_json = json.dumps(fig, cls=PlotlyJSONEncoder)

    # Return the JSON response
    return JSONResponse(content={"chart": chart_json}, media_type="application/json")

@app.post("/data_insights")
async def data_insights(file: UploadFile = File(...)):
    # Read the CSV data
    df = pd.read_csv(file.file)

    # Generate insights
    insights = generate_insights(df)

    # Return the JSON response
    return JSONResponse(content={"insights": insights}, media_type="application/json")

@app.post("/correlation_heatmap")
async def correlation_heatmap(file: UploadFile = File(...)):
    # Read the CSV data
    df = pd.read_csv(file.file)

    # Generate the correlation heatmap
    fig = generate_correlation_matrix(df)

    # Convert the heatmap to JSON
    heatmap_json = json.dumps(fig, cls=PlotlyJSONEncoder)

    # Return the JSON response
    return JSONResponse(content={"heatmap": heatmap_json}, media_type="application/json")

@app.post("/export_to_word")
async def export_to_word(
    x_column: str = Form(...),
    y_column: str = Form(...),
    chart_type: str = Form(...),
    filter_column: Optional[str] = Form(''),
    time_range_start: Optional[str] = Form(''),
    time_range_end: Optional[str] = Form(''),
    file: UploadFile = File(...)
):
    # Read the CSV data
    df = pd.read_csv(file.file)

    # Create a Word document
    doc = Document()

    # Add the dashboard to the document
    doc.add_heading('Dashboard', level=1)
    dashboard_fig = generate_chart(df, x_column, y_column, chart_type)
    dashboard_fig.write_image("dashboard.png")
    doc.add_picture("dashboard.png", width=Inches(6))

    # Add the correlation heatmap to the document
    doc.add_heading('Correlation Heatmap', level=1)
    heatmap_fig = generate_correlation_matrix(df)
    heatmap_fig.write_image("heatmap.png")
    doc.add_picture("heatmap.png", width=Inches(6))

    # Add insights to the document
    doc.add_heading('Data Insights', level=1)
    insights = generate_insights(df)
    dataset_size = insights["dataset_size"]
    doc.add_paragraph(f"Dataset size: {dataset_size}")
    for column, column_insights in insights["column_insights"].items():
        # add each column insight in table form
        doc.add_paragraph(f"Column: {column}")
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Insight'
        hdr_cells[1].text = 'Value'
        for insight, value in column_insights.items():
            row_cells = table.add_row().cells
            row_cells[0].text = insight
            row_cells[1].text = str(value)
        doc.add_paragraph("")

    # Save the document to a file
    doc.save("insights.docx")

    return FileResponse("insights.docx", media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
